"""
POCO 填充引擎单元测试。

测试范围：
  - 单字段替换
  - 多字段替换
  - 表格内替换
  - 页眉替换
  - 页脚替换
  - 缺失字段（占位符不在 mapping 中 → 保持原样）
  - 中文占位符
"""

import os
import struct
import sys
import tempfile
import unittest
import zlib

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Inches, Pt, RGBColor

from poco.filler import fill_template
from poco.scanner import read_docx_text
from poco.extractor import extract_placeholders


# ---- 辅助工具 ---------------------------------------------------------------


def _minimal_png(width, height, r, g, b):
    """生成纯色最小 PNG 字节（用于测试内联图片）。"""

    def _chunk(ctype, data):
        c = ctype + data
        return (
            struct.pack(">I", len(data))
            + c
            + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)
        )

    raw = b""
    for _y in range(height):
        raw += b"\x00"  # filter byte
        for _x in range(width):
            raw += struct.pack("BBB", r, g, b)

    return (
        b"\x89PNG\r\n\x1a\n"
        + _chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0))
        + _chunk(b"IDAT", zlib.compress(raw))
        + _chunk(b"IEND", b"")
    )


class TestFiller(unittest.TestCase):
    """填充引擎测试"""

    def setUp(self):
        """每个测试前创建临时目录"""
        self.tmpdir = tempfile.mkdtemp(prefix="poco_filler_test_")

    def _template_path(self, name="template.docx"):
        return os.path.join(self.tmpdir, name)

    def _output_path(self, name="output.docx"):
        return os.path.join(self.tmpdir, name)

    def _create_template(self, paragraph_texts, path=None):
        """
        创建一个简单 .docx 模板，每个 paragraph_texts 元素是一个段落文本。

        返回模板路径。
        """
        if path is None:
            path = self._template_path()
        doc = Document()
        for text in paragraph_texts:
            doc.add_paragraph(text)
        doc.save(path)
        return path

    def _read_output_text(self, path=None):
        """读取生成的 .docx 文件文本"""
        if path is None:
            path = self._output_path()
        return read_docx_text(path)

    # ------------------------------------------------------------------

    def test_single_placeholder(self):
        """单字段替换"""
        self._create_template(["姓名：{{姓名}}"])
        fill_template(
            self._template_path(),
            self._output_path(),
            {"姓名": "PAN YU"},
        )
        text = self._read_output_text()
        self.assertIn("姓名：PAN YU", text)
        self.assertNotIn("{{姓名}}", text)

    def test_multiple_placeholders(self):
        """多字段替换"""
        self._create_template(["{{姓名}} {{护照号}} {{出生日期}}"])
        fill_template(
            self._template_path(),
            self._output_path(),
            {"姓名": "PAN YU", "护照号": "EJ6376603", "出生日期": "1990-01-15"},
        )
        text = self._read_output_text()
        self.assertIn("PAN YU", text)
        self.assertIn("EJ6376603", text)
        self.assertIn("1990-01-15", text)
        self.assertNotIn("{{", text)
        self.assertNotIn("}}", text)

    def test_table_replacement(self):
        """表格内替换"""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "姓名"
        table.cell(0, 1).text = "{{姓名}}"
        table.cell(1, 0).text = "护照号"
        table.cell(1, 1).text = "{{护照号}}"
        doc.save(self._template_path())

        fill_template(
            self._template_path(),
            self._output_path(),
            {"姓名": "PAN YU", "护照号": "EJ6376603"},
        )
        text = self._read_output_text()
        self.assertIn("PAN YU", text)
        self.assertIn("EJ6376603", text)
        self.assertNotIn("{{", text)

    def test_header_replacement(self):
        """页眉替换"""
        doc = Document()
        doc.add_paragraph("正文内容")

        section = doc.sections[0]
        header = section.header
        header.paragraphs[0].text = "页眉：{{公司名}}"

        doc.save(self._template_path())

        fill_template(
            self._template_path(),
            self._output_path(),
            {"公司名": "ACME Corp"},
        )
        text = self._read_output_text()
        self.assertIn("ACME Corp", text)
        self.assertNotIn("{{公司名}}", text)

    def test_footer_replacement(self):
        """页脚替换"""
        doc = Document()
        doc.add_paragraph("正文内容")

        section = doc.sections[0]
        footer = section.footer
        footer.paragraphs[0].text = "页脚：第{{页码}}页"

        doc.save(self._template_path())

        fill_template(
            self._template_path(),
            self._output_path(),
            {"页码": "1"},
        )
        text = self._read_output_text()
        self.assertIn("第1页", text)
        self.assertNotIn("{{页码}}", text)

    def test_missing_field_raises_error(self):
        """缺失字段 → 抛出 ValueError（rule_engine 必须覆盖全部字段）"""
        self._create_template(["{{姓名}} {{护照号}}"])
        with self.assertRaises(ValueError) as ctx:
            fill_template(
                self._template_path(),
                self._output_path(),
                {"姓名": "PAN YU"},
            )
        self.assertIn("护照号", str(ctx.exception))

    def test_chinese_placeholders(self):
        """中文占位符替换"""
        self._create_template(["客户姓名：{{客户姓名}}，订单号：{{订单号}}"])
        fill_template(
            self._template_path(),
            self._output_path(),
            {"客户姓名": "张三", "订单号": "ORD-001"},
        )
        text = self._read_output_text()
        self.assertIn("客户姓名：张三", text)
        self.assertIn("订单号：ORD-001", text)
        self.assertNotIn("{{", text)

    def test_extra_mapping_keys_ignored(self):
        """mapping 中多余的键不影响输出"""
        self._create_template(["{{姓名}}"])
        fill_template(
            self._template_path(),
            self._output_path(),
            {
                "姓名": "PAN YU",
                "不存在的字段": "IGNORED",  # 模板中没有此占位符
            },
        )
        text = self._read_output_text()
        self.assertIn("PAN YU", text)
        self.assertNotIn("不存在的字段", text)
        self.assertNotIn("IGNORED", text)
        self.assertNotIn("{{", text)

    def test_no_placeholders_passthrough(self):
        """无占位符的模板原样输出"""
        self._create_template(["This is a plain text document."])
        fill_template(
            self._template_path(),
            self._output_path(),
            {"anything": "value"},
        )
        text = self._read_output_text()
        self.assertIn("This is a plain text document.", text)

    def test_empty_mapping_raises_error(self):
        """空 mapping → 抛出 ValueError（所有占位符均缺失）"""
        self._create_template(["{{姓名}} {{护照号}}"])
        with self.assertRaises(ValueError) as ctx:
            fill_template(
                self._template_path(),
                self._output_path(),
                {},
            )
        self.assertIn("姓名", str(ctx.exception))
        self.assertIn("护照号", str(ctx.exception))

    def test_output_directory_auto_created(self):
        """输出目录不存在时自动创建"""
        out_dir = os.path.join(self.tmpdir, "deep", "nested")
        out_path = os.path.join(out_dir, "result.docx")
        self._create_template(["{{姓名}}"])

        self.assertFalse(os.path.exists(out_dir))
        fill_template(
            self._template_path(),
            out_path,
            {"姓名": "PAN YU"},
        )
        self.assertTrue(os.path.exists(out_path))
        text = read_docx_text(out_path)
        self.assertIn("PAN YU", text)

    # ------------------------------------------------------------------
    # 跨 run 占位符 + 图片保留测试（回归：_rebuild_paragraph_runs bug）

    def test_cross_run_placeholder_in_header(self):
        """跨 run 占位符（页眉）：替换后文字正确，run 不被删除"""
        doc = Document()
        doc.add_paragraph("正文内容")

        section = doc.sections[0]
        header = section.header
        para = header.paragraphs[0]

        # 模拟占位符跨两个 run（如 Word 中编辑导致）
        run1 = para.add_run("编号：{{酒店")
        run1.bold = True
        run2 = para.add_run("号公式}}")
        run2.font.size = Pt(10)

        # 记录替换前 run 数量
        run_count_before = len(para.runs)

        doc.save(self._template_path())

        fill_template(
            self._template_path(),
            self._output_path(),
            {"酒店号公式": "26062342824"},
        )

        # 验证输出
        text = self._read_output_text()
        self.assertIn("编号：26062342824", text)
        self.assertNotIn("{{酒店号公式}}", text)
        self.assertNotIn("{{", text)
        self.assertNotIn("}}", text)

        # 验证 run 数量不变（关键：不删除 run）
        out_doc = Document(self._output_path())
        out_header = out_doc.sections[0].header
        out_para = out_header.paragraphs[0]
        self.assertEqual(
            len(out_para.runs), run_count_before,
            "跨 run 替换后 run 数量应保持不变（禁止删除 run）"
        )

    def test_cross_run_placeholder_preserves_image(self):
        """跨 run 占位符 + 图片：替换后图片不丢失"""
        doc = Document()
        doc.add_paragraph("正文内容")

        section = doc.sections[0]
        header = section.header
        para = header.paragraphs[0]

        # 文本 run 1（占位符前半段）
        para.add_run("编号：{{酒店")

        # 生成最小 PNG 并插入图片 run
        img_path = os.path.join(self.tmpdir, "test_pixel.png")
        with open(img_path, "wb") as f:
            f.write(_minimal_png(4, 4, 255, 0, 0))

        para.add_run().add_picture(img_path, width=Inches(0.5))

        # 文本 run 2（占位符后半段）
        para.add_run("号公式}} 结束")

        doc.save(self._template_path())

        fill_template(
            self._template_path(),
            self._output_path(),
            {"酒店号公式": "26062342824"},
        )

        # 验证文本替换正确
        text = self._read_output_text()
        self.assertIn("编号：26062342824 结束", text)
        self.assertNotIn("{{", text)

        # 验证输出文档中图片仍存在
        out_doc = Document(self._output_path())
        out_header = out_doc.sections[0].header
        out_para = out_header.paragraphs[0]

        # 检查 XML 中是否保留 drawing 元素（图片容器）
        from xml.etree import ElementTree as ET
        xml_str = ET.tostring(out_para._element, encoding="unicode")
        # drawing 元素在任意命名空间前缀下都应存在
        self.assertTrue(
            "<drawing" in xml_str or ":drawing" in xml_str,
            "替换后图片必须保留（禁止删除含图片的 run）"
        )

    def test_single_run_with_other_runs_untouched(self):
        """单 run 占位符 + 其他 run（格式各不同）：替换后格式保留、run 数量不变"""
        doc = Document()
        para = doc.add_paragraph()

        # Run 0: 粗体前缀
        r0 = para.add_run("编号：")
        r0.bold = True
        r0.font.size = Pt(14)

        # Run 1: 占位符（普通字体）
        r1 = para.add_run("{{酒店号公式}}")
        r1.font.size = Pt(12)
        r1.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        # Run 2: 后缀（斜体）
        r2 = para.add_run("（必填）")
        r2.italic = True
        r2.font.size = Pt(10)

        run_count_before = len(para.runs)
        doc.save(self._template_path())

        fill_template(
            self._template_path(),
            self._output_path(),
            {"酒店号公式": "26062342824"},
        )

        # 验证文本
        text = self._read_output_text()
        self.assertIn("编号：26062342824（必填）", text)
        self.assertNotIn("{{", text)

        # 验证 run 数量和格式
        out_doc = Document(self._output_path())
        out_para = out_doc.paragraphs[0]
        out_runs = out_para.runs

        self.assertEqual(len(out_runs), run_count_before,
                         "替换后 run 数量应保持不变")

        # Run 0 格式应保留
        self.assertTrue(out_runs[0].bold, "Run 0 粗体应保留")
        self.assertEqual(out_runs[0].text, "编号：")

        # Run 1 格式应保留（仅文字改变）
        self.assertEqual(out_runs[1].text, "26062342824")
        self.assertEqual(out_runs[1].font.size, Pt(12))
        # 颜色保留
        try:
            self.assertEqual(out_runs[1].font.color.rgb,
                             RGBColor(0xFF, 0x00, 0x00))
        except (ValueError, TypeError):
            pass  # 颜色可能以不同方式存储

        # Run 2 应完全不变
        self.assertTrue(out_runs[2].italic, "Run 2 斜体应保留")
        self.assertEqual(out_runs[2].text, "（必填）")

    def test_cross_run_three_runs(self):
        """占位符跨 3 个 run：替换后 run 数量不变"""
        doc = Document()
        para = doc.add_paragraph()

        r0 = para.add_run("开头{{酒店")
        r1 = para.add_run("号公式")
        r2 = para.add_run("}}结尾")

        run_count_before = len(para.runs)
        doc.save(self._template_path())

        fill_template(
            self._template_path(),
            self._output_path(),
            {"酒店号公式": "26062342824"},
        )

        text = self._read_output_text()
        self.assertIn("开头26062342824结尾", text)
        self.assertNotIn("{{", text)

        out_doc = Document(self._output_path())
        out_para = out_doc.paragraphs[0]
        self.assertEqual(len(out_para.runs), run_count_before,
                         "3-run 跨 run 替换后 run 数量应保持不变")


if __name__ == "__main__":
    unittest.main(verbosity=2)
