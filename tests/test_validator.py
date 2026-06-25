"""
POCO 校验与防漏系统单元测试。

测试范围：
  - validate_mapping: 缺失字段 / 格式校验 / 正常通过
  - fill_template 集成: 自动兜底 / 残留报错
  - check_residual: 残留占位符检测
  - register_validator: 自定义校验器
"""

import os
import sys
import tempfile
import unittest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document

from poco.validator import (
    validate_mapping,
    validate_required_fields,
    check_residual,
    register_validator,
    unregister_validator,
    list_validators,
)
from poco.filler import fill_template
from poco.scanner import read_docx_text
from poco.extractor import extract_placeholders


class TestValidateMapping(unittest.TestCase):
    """validate_mapping 单元测试"""

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp(prefix="poco_val_test_")

    def _create_template(self, paragraph_texts, name="template.docx"):
        path = os.path.join(self.tmpdir, name)
        doc = Document()
        for text in paragraph_texts:
            doc.add_paragraph(text)
        doc.save(path)
        return path

    # ------------------------------------------------------------------

    def test_all_fields_present(self):
        """所有字段都在 mapping 中 → 无警告"""
        path = self._create_template(["{{姓名}} {{护照号}}"])
        result = validate_mapping(path, {"姓名": "PAN", "护照号": "EJ6376603"})
        self.assertEqual(result.missing_fields, [])
        self.assertEqual(result.invalid_fields, [])
        self.assertFalse(result.has_warnings)

    def test_missing_field_detected(self):
        """模板有但 mapping 没有的字段 → missing_fields"""
        path = self._create_template(["{{姓名}} {{护照号}} {{出生日期}}"])
        result = validate_mapping(path, {"姓名": "PAN"})
        self.assertIn("护照号", result.missing_fields)
        self.assertIn("出生日期", result.missing_fields)
        self.assertTrue(result.has_warnings)

    def test_passport_valid(self):
        """有效护照号 → 无格式警告"""
        path = self._create_template(["{{护照号}}"])
        result = validate_mapping(path, {"护照号": "EJ6376603"})
        self.assertEqual(result.invalid_fields, [])

    def test_passport_too_short(self):
        """护照号长度 < 6 → invalid_fields"""
        path = self._create_template(["{{护照号}}"])
        result = validate_mapping(path, {"护照号": "EJ6"})
        self.assertEqual(len(result.invalid_fields), 1)
        self.assertIn("长度不足", result.invalid_fields[0]["reason"])

    def test_passport_has_space(self):
        """护照号含空格 → invalid_fields"""
        path = self._create_template(["{{护照号}}"])
        result = validate_mapping(path, {"护照号": "EJ 637"})
        self.assertEqual(len(result.invalid_fields), 1)
        self.assertIn("空格", result.invalid_fields[0]["reason"])

    def test_passport_special_chars(self):
        """护照号含特殊字符 → invalid_fields"""
        path = self._create_template(["{{护照号}}"])
        result = validate_mapping(path, {"护照号": "ABC-123"})
        self.assertEqual(len(result.invalid_fields), 1)
        self.assertIn("字母和数字", result.invalid_fields[0]["reason"])

    def test_report_clean(self):
        """校验通过时 report 包含 [OK]"""
        path = self._create_template(["{{姓名}}"])
        result = validate_mapping(path, {"姓名": "PAN"})
        report = result.report()
        self.assertIn("[OK]", report)

    def test_report_with_warnings(self):
        """校验警告时 report 包含 [WARNING]"""
        path = self._create_template(["{{姓名}} {{护照号}}"])
        result = validate_mapping(path, {"姓名": "PAN"})
        report = result.report()
        self.assertIn("[WARNING]", report)
        self.assertIn("护照号", report)

    def test_summary_structure(self):
        """summary() 返回结构化字典"""
        path = self._create_template(["{{姓名}} {{护照号}}"])
        result = validate_mapping(path, {"护照号": "E J"})
        s = result.summary()
        self.assertIsInstance(s, dict)
        self.assertIn("missing_fields", s)
        self.assertIn("invalid_fields", s)
        self.assertIn("replaced_fallback_fields", s)
        self.assertIn("姓名", s["missing_fields"])

    def test_custom_validator(self):
        """注册自定义校验器 → 生效"""
        path = self._create_template(["{{订单号}}"])

        def _validate_order(value):
            if not value.startswith("ORD-"):
                return "订单号必须以 ORD- 开头"
            return None

        register_validator("订单号", _validate_order)
        try:
            result = validate_mapping(path, {"订单号": "BAD"})
            self.assertEqual(len(result.invalid_fields), 1)
            self.assertIn("ORD-", result.invalid_fields[0]["reason"])

            result2 = validate_mapping(path, {"订单号": "ORD-001"})
            self.assertEqual(len(result2.invalid_fields), 0)
        finally:
            unregister_validator("订单号")


class TestFillTemplateWithValidation(unittest.TestCase):
    """fill_template 校验集成测试"""

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp(prefix="poco_val_int_")

    def _template_path(self, name="template.docx"):
        return os.path.join(self.tmpdir, name)

    def _output_path(self, name="output.docx"):
        return os.path.join(self.tmpdir, name)

    def _create_template(self, paragraph_texts, path=None):
        if path is None:
            path = self._template_path()
        doc = Document()
        for text in paragraph_texts:
            doc.add_paragraph(text)
        doc.save(path)
        return path

    # ------------------------------------------------------------------

    def test_missing_field_raises_error(self):
        """
        缺失字段 → 抛出 ValueError（rule_engine 必须覆盖全部字段）
        """
        self._create_template(["{{姓名}} {{护照号}}"])
        with self.assertRaises(ValueError) as ctx:
            fill_template(
                self._template_path(),
                self._output_path(),
                {"姓名": "PAN"},
            )
        self.assertIn("护照号", str(ctx.exception))

    def test_invalid_passport_still_fills(self):
        """
        护照号格式错误 → 打印 warning，但仍正常填充
        """
        self._create_template(["{{护照号}}"])
        fill_template(
            self._template_path(),
            self._output_path(),
            {"护照号": "E J"},
        )
        text = read_docx_text(self._output_path())
        # 值仍然被填入了文档（warning 不阻断填充）
        self.assertIn("E J", text)
        self.assertNotIn("{{", text)

    def test_all_fields_valid_clean_output(self):
        """全部字段合法 → 无警告，输出干净"""
        self._create_template(["{{姓名}} {{护照号}}"])
        fill_template(
            self._template_path(),
            self._output_path(),
            {"姓名": "PAN YU", "护照号": "EJ6376603"},
        )
        text = read_docx_text(self._output_path())
        self.assertIn("PAN YU", text)
        self.assertIn("EJ6376603", text)
        self.assertNotIn("[未填写", text)
        self.assertNotIn("{{", text)

    def test_table_with_missing_field_raises(self):
        """表格内缺失字段 → 抛出 ValueError"""
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "护照号"
        table.cell(0, 1).text = "{{护照号}}"
        doc.save(self._template_path())

        with self.assertRaises(ValueError) as ctx:
            fill_template(
                self._template_path(),
                self._output_path(),
                {},
            )
        self.assertIn("护照号", str(ctx.exception))

    def test_chinese_missing_field_raises(self):
        """中文占位符缺失 → 抛出 ValueError"""
        self._create_template(["客户：{{客户姓名}} 订单：{{订单号}}"])
        with self.assertRaises(ValueError) as ctx:
            fill_template(
                self._template_path(),
                self._output_path(),
                {},
            )
        self.assertIn("客户姓名", str(ctx.exception))
        self.assertIn("订单号", str(ctx.exception))


class TestCheckResidual(unittest.TestCase):
    """check_residual 检测测试"""

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp(prefix="poco_val_res_")

    def test_no_residual(self):
        """无占位符文档 → 返回空列表"""
        path = os.path.join(self.tmpdir, "clean.docx")
        doc = Document()
        doc.add_paragraph("纯文本，无占位符。")
        doc.save(path)

        residual = check_residual(path)
        self.assertEqual(residual, [])

    def test_has_residual(self):
        """有占位符 → 返回占位符列表"""
        path = os.path.join(self.tmpdir, "dirty.docx")
        doc = Document()
        doc.add_paragraph("{{未处理字段}}")
        doc.save(path)

        residual = check_residual(path)
        self.assertIn("未处理字段", residual)

    def test_fill_template_raises_on_residual(self):
        """
        fill_template 内部检测到残留 → raise ValueError

        注：正常情况下不会发生（因验证阶段已拦截缺失字段）。
        此测试验证检测逻辑存在。
        """
        # 直接构建一个包含 {{xxx}} 的 docx 并验证 check_residual 能发现
        path = os.path.join(self.tmpdir, "bad.docx")
        doc = Document()
        doc.add_paragraph("{{漏网之鱼}}")
        doc.save(path)

        residual = check_residual(path)
        self.assertGreater(len(residual), 0)
        # fill_template 流程会在阶段 3 调用 check_residual 并 raise
        # 这里验证 check_residual 本身能正确检测


class TestValidateRequiredFields(unittest.TestCase):
    """validate_required_fields 单元测试"""

    def test_all_fields_filled(self):
        """所有必填字段已填写 → 返回空列表"""
        user_info = {
            "姓": "PAN",
            "名": "YU",
            "护照号": "EJ6376603",
            "到达日期": "2026-06-23",
        }
        self.assertEqual(validate_required_fields(user_info), [])

    def test_empty_last_name(self):
        """姓为空 → 返回['姓 未填写']"""
        user_info = {
            "姓": "",
            "名": "YU",
            "护照号": "EJ6376603",
            "到达日期": "2026-06-23",
        }
        errors = validate_required_fields(user_info)
        self.assertIn("姓 未填写", errors)
        self.assertNotIn("名 未填写", errors)

    def test_empty_first_name(self):
        """名为空 → 返回['名 未填写']"""
        user_info = {
            "姓": "PAN",
            "名": "",
            "护照号": "EJ6376603",
            "到达日期": "2026-06-23",
        }
        errors = validate_required_fields(user_info)
        self.assertIn("名 未填写", errors)

    def test_whitespace_only(self):
        """仅空格 → 视为未填写"""
        user_info = {
            "姓": "PAN",
            "名": "   ",
            "护照号": "EJ6376603",
            "到达日期": "2026-06-23",
        }
        errors = validate_required_fields(user_info)
        self.assertIn("名 未填写", errors)

    def test_trim_empty(self):
        """trim后为空 → 视为未填写"""
        user_info = {
            "姓": "PAN",
            "名": "  \t  ",
            "护照号": "\n ",
            "到达日期": "2026-06-23",
        }
        errors = validate_required_fields(user_info)
        self.assertIn("名 未填写", errors)
        self.assertIn("护照号 未填写", errors)

    def test_multiple_missing(self):
        """同时缺少多个字段 → 全部列出"""
        user_info = {
            "姓": "",
            "名": "",
            "护照号": "",
            "到达日期": "2026-06-23",
        }
        errors = validate_required_fields(user_info)
        self.assertEqual(len(errors), 3)
        self.assertIn("姓 未填写", errors)
        self.assertIn("名 未填写", errors)
        self.assertIn("护照号 未填写", errors)

    def test_all_empty(self):
        """全部为空 → 返回全部 4 个字段错误"""
        user_info = {
            "姓": "",
            "名": "",
            "护照号": "",
            "到达日期": "",
        }
        errors = validate_required_fields(user_info)
        self.assertEqual(len(errors), 4)

    def test_missing_keys(self):
        """user_info 中缺少字段 key → 视为未填写"""
        user_info = {
            "姓": "PAN",
        }
        errors = validate_required_fields(user_info)
        self.assertIn("名 未填写", errors)
        self.assertIn("护照号 未填写", errors)
        self.assertIn("到达日期 未填写", errors)

    def test_none_value(self):
        """字段值为 None → 视为未填写"""
        user_info = {
            "姓": "PAN",
            "名": None,
            "护照号": "EJ6376603",
            "到达日期": "2026-06-23",
        }
        errors = validate_required_fields(user_info)
        self.assertIn("名 未填写", errors)


class TestValidatorRegistry(unittest.TestCase):
    """校验器注册表测试"""

    def test_list_validators_includes_builtins(self):
        """内置护照号、电话号校验器已注册"""
        vlist = list_validators()
        self.assertIn("护照号", vlist)
        self.assertIn("电话号", vlist)

    def test_unregister_removes_validator(self):
        """unregister 后不再校验"""
        register_validator("测试字段", lambda v: None)
        self.assertIn("测试字段", list_validators())
        unregister_validator("测试字段")
        self.assertNotIn("测试字段", list_validators())


if __name__ == "__main__":
    unittest.main(verbosity=2)
