"""
POCO 多格式架构测试。

测试范围:
  - FormatDetector:  文件格式识别（docx / odt / unknown）
  - BaseRenderer:    ABC 接口契约
  - DocumentEngine:  统一引擎（自动检测 + 调度）
  - DocxRenderer:    DOCX 渲染器（兼容旧行为）
  - OdtRenderer:     ODT 渲染器（文本 / 表格 / 跨元素占位符）
  - scan_template_text: 格式无关的文本扫描
  - 统一占位符系统（docx 和 odt 使用相同的 {{xxx}} 格式）
"""

import os
import struct
import sys
import tempfile
import unittest
import zlib
import zipfile
import io

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from poco.core.format_detector import detect_format, is_supported_format, SUPPORTED_FORMATS
from poco.core.base_renderer import BaseRenderer
from poco.core.engine import (
    DocumentEngine,
    scan_template_text,
    register_renderer,
    get_renderer,
    list_registered_formats,
)
from poco.renderers.docx_renderer import DocxRenderer
from poco.renderers.odt_renderer import OdtRenderer


# ---- 辅助工具 ---------------------------------------------------------------


def _minimal_png(width, height, r, g, b):
    """生成纯色最小 PNG 字节"""
    def _chunk(ctype, data):
        c = ctype + data
        return (
            struct.pack(">I", len(data))
            + c
            + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)
        )
    raw = b""
    for _y in range(height):
        raw += b"\x00"
        for _x in range(width):
            raw += struct.pack("BBB", r, g, b)
    return (
        b"\x89PNG\r\n\x1a\n"
        + _chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0))
        + _chunk(b"IDAT", zlib.compress(raw))
        + _chunk(b"IEND", b"")
    )


def _create_minimal_docx(path: str, paragraphs: list) -> str:
    """创建最小 .docx 文件用于测试"""
    from docx import Document
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(path)
    return path


def _create_minimal_odt(path: str, paragraphs: list) -> str:
    """
    创建最小 .odt 文件用于测试。

    直接构建 ODF ZIP 归档，包含:
      - mimetype
      - content.xml (带占位符文本的段落)
      - META-INF/manifest.xml
    """
    # content.xml 模板
    content_paragraphs = ""
    for text in paragraphs:
        content_paragraphs += (
            f'    <text:p text:style-name="Standard">{text}</text:p>\n'
        )

    content_xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<office:document-content
 xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"
 xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0"
 xmlns:table="urn:oasis:names:tc:opendocument:xmlns:table:1.0"
 xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0"
 office:version="1.2">
  <office:body>
    <office:text>
{content_paragraphs}
    </office:text>
  </office:body>
</office:document-content>"""

    manifest_xml = """<?xml version="1.0" encoding="UTF-8"?>
<manifest:manifest
 xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0"
 manifest:version="1.2">
  <manifest:file-entry
   manifest:full-path="/"
   manifest:version="1.2"
   manifest:media-type="application/vnd.oasis.opendocument.text"/>
  <manifest:file-entry
   manifest:full-path="content.xml"
   manifest:media-type="text/xml"/>
</manifest:manifest>"""

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("mimetype", "application/vnd.oasis.opendocument.text")
        zf.writestr("content.xml", content_xml.encode("utf-8"))
        zf.writestr("META-INF/manifest.xml", manifest_xml.encode("utf-8"))

    return path


# =============================================================================
#  Format Detector Tests
# =============================================================================


class TestFormatDetector(unittest.TestCase):
    """格式检测器测试"""

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp(prefix="poco_fmt_test_")

    def _path(self, name):
        return os.path.join(self.tmpdir, name)

    def test_detect_docx_by_extension(self):
        """通过扩展名 + ZIP 签名识别 docx"""
        path = self._path("test.docx")
        _create_minimal_docx(path, ["Hello {{name}}"])
        fmt = detect_format(path)
        self.assertEqual(fmt, "docx")

    def test_detect_odt_by_extension(self):
        """通过扩展名 + ZIP 签名识别 odt"""
        path = self._path("test.odt")
        _create_minimal_odt(path, ["Hello {{name}}"])
        fmt = detect_format(path)
        self.assertEqual(fmt, "odt")

    def test_detect_unknown_for_pdf(self):
        """非支持格式返回 unknown"""
        path = self._path("test.pdf")
        with open(path, "w") as f:
            f.write("not a zip file")
        fmt = detect_format(path)
        self.assertEqual(fmt, "unknown")

    def test_detect_unknown_for_nonexistent(self):
        """不存在的文件返回 unknown"""
        fmt = detect_format("/nonexistent/file.docx")
        self.assertEqual(fmt, "unknown")

    def test_detect_docx_by_content_when_wrong_extension(self):
        """扩展名错误但内容为 docx → 仍识别为 docx"""
        path = self._path("test.wrong")
        _create_minimal_docx(path, ["Hello {{name}}"])
        fmt = detect_format(path)
        self.assertEqual(fmt, "docx")

    def test_detect_odt_by_content_when_wrong_extension(self):
        """扩展名错误但内容为 odt → 仍识别为 odt"""
        path = self._path("test.wrong")
        _create_minimal_odt(path, ["Hello {{name}}"])
        fmt = detect_format(path)
        self.assertEqual(fmt, "odt")

    def test_is_supported_format(self):
        """is_supported_format 正确判断"""
        docx_path = self._path("t.docx")
        _create_minimal_docx(docx_path, ["test"])
        self.assertTrue(is_supported_format(docx_path))

        odt_path = self._path("t.odt")
        _create_minimal_odt(odt_path, ["test"])
        self.assertTrue(is_supported_format(odt_path))

        self.assertFalse(is_supported_format(self._path("t.pdf")))

    def test_supported_formats_list(self):
        """SUPPORTED_FORMATS 包含 docx 和 odt"""
        self.assertIn("docx", SUPPORTED_FORMATS)
        self.assertIn("odt", SUPPORTED_FORMATS)

    def test_case_insensitive_extension(self):
        """扩展名大小写不敏感"""
        path = self._path("TEST.DOCX")
        _create_minimal_docx(path, ["Hello"])
        fmt = detect_format(path)
        self.assertEqual(fmt, "docx")

        path2 = self._path("TEST.ODT")
        _create_minimal_odt(path2, ["Hello"])
        fmt2 = detect_format(path2)
        self.assertEqual(fmt2, "odt")


# =============================================================================
#  BaseRenderer Tests
# =============================================================================


class TestBaseRenderer(unittest.TestCase):
    """BaseRenderer ABC 测试"""

    def test_cannot_instantiate_abstract(self):
        """ABC 不能直接实例化"""
        with self.assertRaises(TypeError):
            BaseRenderer()

    def test_concrete_subclass_instantiates(self):
        """具体子类可以实例化"""

        class ConcreteRenderer(BaseRenderer):
            def load(self, path): pass
            def fill(self, data): pass
            def save(self, path): pass

        renderer = ConcreteRenderer()
        self.assertIsInstance(renderer, BaseRenderer)

    def test_default_extract_placeholders_empty(self):
        """默认 extract_placeholders 返回空列表"""

        class ConcreteRenderer(BaseRenderer):
            def load(self, path): pass
            def fill(self, data): pass
            def save(self, path): pass

        renderer = ConcreteRenderer()
        self.assertEqual(renderer.extract_placeholders(), [])

    def test_default_validate_returns_missing(self):
        """默认 validate 基于 extract_placeholders 检测缺失字段"""

        class ConcreteRenderer(BaseRenderer):
            def load(self, path): pass
            def fill(self, data): pass
            def save(self, path): pass

            def extract_placeholders(self):
                return ["name", "date"]

        renderer = ConcreteRenderer()
        missing = renderer.validate({"name": "PAN"})
        self.assertEqual(missing, ["date"])

    def test_error_tag_formatting(self):
        """错误标签格式化正确"""
        tag = BaseRenderer._error_tag("TEST", "something wrong")
        self.assertEqual(tag, "[TEST] something wrong")


# =============================================================================
#  DocxRenderer Tests（兼容旧 filler 行为）
# =============================================================================


class TestDocxRenderer(unittest.TestCase):
    """DocxRenderer 渲染器测试"""

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp(prefix="poco_docxr_test_")

    def _path(self, name):
        return os.path.join(self.tmpdir, name)

    def test_basic_process(self):
        """基本 load → fill → save 流程"""
        tpl = self._path("template.docx")
        out = self._path("output.docx")
        _create_minimal_docx(tpl, ["姓名：{{姓名}}"])

        renderer = DocxRenderer()
        renderer.load(tpl)
        renderer.fill({"姓名": "PAN YU"})
        renderer.save(out)

        self.assertTrue(os.path.exists(out))
        from poco.scanner import read_docx_text
        text = read_docx_text(out)
        self.assertIn("PAN YU", text)
        self.assertNotIn("{{", text)

    def test_one_shot_process(self):
        """一站式 process 方法"""
        tpl = self._path("template.docx")
        out = self._path("output.docx")
        _create_minimal_docx(tpl, ["{{greeting}} World"])

        renderer = DocxRenderer()
        renderer.process(tpl, {"greeting": "Hello"}, out)

        from poco.scanner import read_docx_text
        text = read_docx_text(out)
        self.assertIn("Hello World", text)

    def test_extract_placeholders(self):
        """占位符提取"""
        tpl = self._path("template.docx")
        _create_minimal_docx(tpl, ["{{name}} {{date}} {{name}}"])

        renderer = DocxRenderer()
        renderer.load(tpl)
        ph = renderer.extract_placeholders()
        self.assertEqual(ph, ["name", "date"])

    def test_validate_returns_missing(self):
        """校验返回缺失字段"""
        tpl = self._path("template.docx")
        _create_minimal_docx(tpl, ["{{name}} {{date}}"])

        renderer = DocxRenderer()
        renderer.load(tpl)
        missing = renderer.validate({"name": "PAN"})
        self.assertIn("date", missing)

    def test_load_nonexistent_raises(self):
        """加载不存在的文件抛 FileNotFoundError"""
        renderer = DocxRenderer()
        with self.assertRaises(FileNotFoundError):
            renderer.load("/nonexistent/template.docx")

    def test_fill_without_load_raises(self):
        """未加载时调用 fill 抛 RuntimeError"""
        renderer = DocxRenderer()
        with self.assertRaises(RuntimeError):
            renderer.fill({"name": "test"})

    def test_save_without_load_raises(self):
        """未加载时调用 save 抛 RuntimeError"""
        renderer = DocxRenderer()
        with self.assertRaises(RuntimeError):
            renderer.save("/tmp/out.docx")

    def test_table_replacement(self):
        """表格内占位符替换"""
        from docx import Document

        tpl = self._path("template.docx")
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "姓名"
        table.cell(0, 1).text = "{{姓名}}"
        table.cell(1, 0).text = "护照号"
        table.cell(1, 1).text = "{{护照号}}"
        doc.save(tpl)

        out = self._path("output.docx")
        renderer = DocxRenderer()
        renderer.process(tpl, {"姓名": "PAN YU", "护照号": "EJ6376603"}, out)

        from poco.scanner import read_docx_text
        text = read_docx_text(out)
        self.assertIn("PAN YU", text)
        self.assertIn("EJ6376603", text)

    def test_chinese_placeholders(self):
        """中文占位符替换"""
        tpl = self._path("template.docx")
        out = self._path("output.docx")
        _create_minimal_docx(tpl, ["客户姓名：{{客户姓名}}，订单号：{{订单号}}"])

        renderer = DocxRenderer()
        renderer.process(tpl, {"客户姓名": "张三", "订单号": "ORD-001"}, out)

        from poco.scanner import read_docx_text
        text = read_docx_text(out)
        self.assertIn("张三", text)
        self.assertIn("ORD-001", text)

    def test_auto_create_output_dir(self):
        """输出目录不存在时自动创建"""
        tpl = self._path("template.docx")
        out_dir = os.path.join(self.tmpdir, "deep", "nested")
        out = os.path.join(out_dir, "result.docx")
        _create_minimal_docx(tpl, ["{{name}}"])

        renderer = DocxRenderer()
        renderer.process(tpl, {"name": "PAN"}, out)
        self.assertTrue(os.path.exists(out))


# =============================================================================
#  OdtRenderer Tests
# =============================================================================


class TestOdtRenderer(unittest.TestCase):
    """OdtRenderer 渲染器测试"""

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp(prefix="poco_odtr_test_")

    def _path(self, name):
        return os.path.join(self.tmpdir, name)

    def _read_odt_text(self, path):
        """读取 ODT 文件中的文本"""
        with zipfile.ZipFile(path, "r") as zf:
            content = zf.read("content.xml").decode("utf-8")
        return content

    def test_basic_process(self):
        """基本 load → fill → save 流程"""
        tpl = self._path("template.odt")
        out = self._path("output.odt")
        _create_minimal_odt(tpl, ["姓名：{{姓名}}"])

        renderer = OdtRenderer()
        renderer.load(tpl)
        renderer.fill({"姓名": "PAN YU"})
        renderer.save(out)

        self.assertTrue(os.path.exists(out))
        content = self._read_odt_text(out)
        self.assertIn("PAN YU", content)
        self.assertNotIn("{{姓名}}", content)

    def test_one_shot_process(self):
        """一站式 process 方法"""
        tpl = self._path("template.odt")
        out = self._path("output.odt")
        _create_minimal_odt(tpl, ["{{greeting}} World"])

        renderer = OdtRenderer()
        renderer.process(tpl, {"greeting": "Hello"}, out)

        content = self._read_odt_text(out)
        self.assertIn("Hello World", content)

    def test_multiple_placeholders(self):
        """多占位符替换"""
        tpl = self._path("template.odt")
        out = self._path("output.odt")
        _create_minimal_odt(tpl, [
            "姓名：{{姓名}}",
            "护照号：{{护照号}}",
            "日期：{{日期}}",
        ])

        renderer = OdtRenderer()
        renderer.process(
            tpl,
            {"姓名": "PAN YU", "护照号": "EJ6376603", "日期": "2026-06-27"},
            out,
        )

        content = self._read_odt_text(out)
        self.assertIn("PAN YU", content)
        self.assertIn("EJ6376603", content)
        self.assertIn("2026-06-27", content)

    def test_extract_placeholders(self):
        """ODT 占位符提取"""
        tpl = self._path("template.odt")
        _create_minimal_odt(tpl, ["{{name}} {{date}} {{name}}"])

        renderer = OdtRenderer()
        renderer.load(tpl)
        ph = renderer.extract_placeholders()
        self.assertEqual(ph, ["name", "date"])

    def test_validate_returns_missing(self):
        """校验返回缺失字段"""
        tpl = self._path("template.odt")
        _create_minimal_odt(tpl, ["{{name}} {{date}}"])

        renderer = OdtRenderer()
        renderer.load(tpl)
        missing = renderer.validate({"name": "PAN"})
        self.assertIn("date", missing)

    def test_chinese_placeholders(self):
        """中文占位符替换"""
        tpl = self._path("template.odt")
        out = self._path("output.odt")
        _create_minimal_odt(tpl, ["客户姓名：{{客户姓名}}，订单号：{{订单号}}"])

        renderer = OdtRenderer()
        renderer.process(
            tpl,
            {"客户姓名": "张三", "订单号": "ORD-001"},
            out,
        )

        content = self._read_odt_text(out)
        self.assertIn("张三", content)
        self.assertIn("ORD-001", content)
        self.assertNotIn("{{", content)

    def test_load_nonexistent_raises(self):
        """加载不存在的文件抛 FileNotFoundError"""
        renderer = OdtRenderer()
        with self.assertRaises(FileNotFoundError):
            renderer.load("/nonexistent/template.odt")

    def test_load_invalid_odt_raises(self):
        """加载无效 ODT 抛 ValueError"""
        path = self._path("bad.odt")
        with open(path, "w") as f:
            f.write("not a zip file")
        renderer = OdtRenderer()
        with self.assertRaises(ValueError):
            renderer.load(path)

    def test_fill_without_load_raises(self):
        """未加载时调用 fill 抛 RuntimeError"""
        renderer = OdtRenderer()
        with self.assertRaises(RuntimeError):
            renderer.fill({"name": "test"})

    def test_save_without_load_raises(self):
        """未加载时调用 save 抛 RuntimeError"""
        renderer = OdtRenderer()
        with self.assertRaises(RuntimeError):
            renderer.save("/tmp/out.odt")

    def test_auto_create_output_dir(self):
        """输出目录不存在时自动创建"""
        tpl = self._path("template.odt")
        out_dir = os.path.join(self.tmpdir, "deep", "nested")
        out = os.path.join(out_dir, "result.odt")
        _create_minimal_odt(tpl, ["{{name}}"])

        renderer = OdtRenderer()
        renderer.process(tpl, {"name": "PAN"}, out)
        self.assertTrue(os.path.exists(out))

    def test_no_placeholders_passthrough(self):
        """无占位符的模板原样输出"""
        tpl = self._path("template.odt")
        out = self._path("output.odt")
        _create_minimal_odt(tpl, ["This is plain text."])

        renderer = OdtRenderer()
        renderer.process(tpl, {"anything": "value"}, out)

        content = self._read_odt_text(out)
        self.assertIn("This is plain text.", content)

    def test_mimetype_preserved(self):
        """mimetype 在保存后保持不变"""
        tpl = self._path("template.odt")
        out = self._path("output.odt")
        _create_minimal_odt(tpl, ["{{name}}"])

        renderer = OdtRenderer()
        renderer.process(tpl, {"name": "PAN"}, out)

        with zipfile.ZipFile(out, "r") as zf:
            mime = zf.read("mimetype").decode("utf-8").strip()
        self.assertIn("opendocument.text", mime)

    def test_odt_with_spans(self):
        """包含 text:span 的 ODT 占位符替换"""
        tpl = self._path("template.odt")
        out = self._path("output.odt")

        # 构建带 span 的 content.xml
        content_xml = """<?xml version="1.0" encoding="UTF-8"?>
<office:document-content
 xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"
 xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0"
 xmlns:table="urn:oasis:names:tc:opendocument:xmlns:table:1.0"
 office:version="1.2">
  <office:body>
    <office:text>
      <text:p text:style-name="Standard">
        <text:span text:style-name="T1">编号：{{酒店</text:span>
        <text:span text:style-name="T2">号公式}}</text:span>
      </text:p>
    </office:text>
  </office:body>
</office:document-content>"""

        manifest_xml = """<?xml version="1.0" encoding="UTF-8"?>
<manifest:manifest
 xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0"
 manifest:version="1.2">
  <manifest:file-entry manifest:full-path="/"
   manifest:media-type="application/vnd.oasis.opendocument.text"/>
  <manifest:file-entry manifest:full-path="content.xml"
   manifest:media-type="text/xml"/>
</manifest:manifest>"""

        with zipfile.ZipFile(tpl, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("mimetype", "application/vnd.oasis.opendocument.text")
            zf.writestr("content.xml", content_xml.encode("utf-8"))
            zf.writestr("META-INF/manifest.xml", manifest_xml.encode("utf-8"))

        renderer = OdtRenderer()
        renderer.process(tpl, {"酒店号公式": "26062342824"}, out)

        content = self._read_odt_text(out)
        self.assertIn("编号：26062342824", content)
        self.assertNotIn("{{酒店", content)
        self.assertNotIn("号公式}}", content)


# =============================================================================
#  DocumentEngine Tests
# =============================================================================


class TestDocumentEngine(unittest.TestCase):
    """DocumentEngine 统一引擎测试"""

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp(prefix="poco_eng_test_")

    def _path(self, name):
        return os.path.join(self.tmpdir, name)

    def test_process_docx(self):
        """引擎自动检测 docx 并处理"""
        tpl = self._path("template.docx")
        out = self._path("output.docx")
        _create_minimal_docx(tpl, ["Hello {{name}}"])

        engine = DocumentEngine()
        fmt = engine.process(tpl, {"name": "World"}, out)

        self.assertEqual(fmt, "docx")
        from poco.scanner import read_docx_text
        text = read_docx_text(out)
        self.assertIn("Hello World", text)

    def test_process_odt(self):
        """引擎自动检测 odt 并处理"""
        tpl = self._path("template.odt")
        out = self._path("output.odt")
        _create_minimal_odt(tpl, ["Hello {{name}}"])

        engine = DocumentEngine()
        fmt = engine.process(tpl, {"name": "World"}, out)

        self.assertEqual(fmt, "odt")
        with zipfile.ZipFile(out, "r") as zf:
            content = zf.read("content.xml").decode("utf-8")
        self.assertIn("Hello World", content)

    def test_process_unknown_format_raises(self):
        """不支持的格式抛 ValueError"""
        path = self._path("test.pdf")
        with open(path, "w") as f:
            f.write("not a document")

        engine = DocumentEngine()
        with self.assertRaises(ValueError):
            engine.process(path, {"x": "y"}, self._path("out.pdf"))

    def test_process_batch(self):
        """批量处理多格式文档"""
        tpl1 = self._path("t1.docx")
        tpl2 = self._path("t2.odt")
        out1 = self._path("out1.docx")
        out2 = self._path("out2.odt")

        _create_minimal_docx(tpl1, ["{{a}}"])
        _create_minimal_odt(tpl2, ["{{b}}"])

        engine = DocumentEngine()
        results = engine.process_batch([
            (tpl1, {"a": "AAA"}, out1),
            (tpl2, {"b": "BBB"}, out2),
        ])

        self.assertEqual(len(results), 2)
        self.assertIsNone(results[0]["error"])
        self.assertIsNone(results[1]["error"])
        self.assertEqual(results[0]["format"], "docx")
        self.assertEqual(results[1]["format"], "odt")

    def test_process_batch_individual_failure(self):
        """批量处理中单个失败不影响其他"""
        tpl1 = self._path("t1.docx")
        out1 = self._path("out1.docx")
        _create_minimal_docx(tpl1, ["{{a}}"])

        engine = DocumentEngine()
        results = engine.process_batch([
            (tpl1, {"a": "AAA"}, out1),
            ("/nonexistent.docx", {}, self._path("out2.docx")),
        ])

        self.assertEqual(len(results), 2)
        self.assertIsNone(results[0]["error"])
        self.assertIsNotNone(results[1]["error"])

    def test_preview_placeholders_docx(self):
        """预览 docx 占位符"""
        tpl = self._path("template.docx")
        _create_minimal_docx(tpl, ["{{name}} {{date}}"])

        engine = DocumentEngine()
        ph = engine.preview_placeholders(tpl)
        self.assertEqual(ph, ["name", "date"])

    def test_preview_placeholders_odt(self):
        """预览 odt 占位符"""
        tpl = self._path("template.odt")
        _create_minimal_odt(tpl, ["{{name}} {{date}}"])

        engine = DocumentEngine()
        ph = engine.preview_placeholders(tpl)
        self.assertEqual(ph, ["name", "date"])

    def test_supported_formats_property(self):
        """supported_formats 属性返回已注册格式"""
        engine = DocumentEngine()
        formats = engine.supported_formats
        self.assertIn("docx", formats)
        self.assertIn("odt", formats)

    def test_explicit_renderer_class(self):
        """显式指定渲染器类跳过格式检测"""
        tpl = self._path("template.docx")
        out = self._path("output.docx")
        _create_minimal_docx(tpl, ["{{x}}"])

        engine = DocumentEngine()
        fmt = engine.process(tpl, {"x": "Y"}, out, renderer_class=DocxRenderer)
        self.assertEqual(fmt, "docx")


# =============================================================================
#  Unified Placeholder System Tests（占位符统一性）
# =============================================================================


class TestUnifiedPlaceholderSystem(unittest.TestCase):
    """占位符统一性测试：docx 和 odt 使用相同的 {{xxx}} 格式"""

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp(prefix="poco_uph_test_")

    def _path(self, name):
        return os.path.join(self.tmpdir, name)

    def test_same_mapping_works_for_both_formats(self):
        """相同的 mapping 同时用于 docx 和 odt"""
        mapping = {"姓名": "PAN YU", "护照号": "EJ6376603", "日期": "2026-06-27"}

        # DOCX
        docx_tpl = self._path("t.docx")
        docx_out = self._path("out.docx")
        _create_minimal_docx(docx_tpl, [
            "姓名：{{姓名}}",
            "护照号：{{护照号}}",
            "日期：{{日期}}",
        ])
        DocxRenderer().process(docx_tpl, mapping, docx_out)

        # ODT
        odt_tpl = self._path("t.odt")
        odt_out = self._path("out.odt")
        _create_minimal_odt(odt_tpl, [
            "姓名：{{姓名}}",
            "护照号：{{护照号}}",
            "日期：{{日期}}",
        ])
        OdtRenderer().process(odt_tpl, mapping, odt_out)

        # 验证 docx 输出
        from poco.scanner import read_docx_text
        docx_text = read_docx_text(docx_out)
        self.assertIn("PAN YU", docx_text)
        self.assertIn("EJ6376603", docx_text)

        # 验证 odt 输出
        with zipfile.ZipFile(odt_out, "r") as zf:
            odt_content = zf.read("content.xml").decode("utf-8")
        self.assertIn("PAN YU", odt_content)
        self.assertIn("EJ6376603", odt_content)

    def test_placeholder_format_identical(self):
        """占位符格式 docx 和 odt 完全一致"""
        from poco.extractor import extract_placeholders

        # DOCX 风格的占位符文本
        text = "{{姓名}} {{护照号}} {{随机四位数}}"
        ph = extract_placeholders(text)
        self.assertEqual(ph, ["姓名", "护照号", "随机四位数"])

        # 相同的文本在 ODT 中也应提取出相同的占位符
        # (extractor 是格式无关的)
        self.assertEqual(len(ph), 3)

    def test_scan_template_text_docx(self):
        """scan_template_text 从 docx 提取文本"""
        tpl = self._path("t.docx")
        _create_minimal_docx(tpl, ["Hello {{name}}"])

        text = scan_template_text(tpl)
        self.assertIn("Hello {{name}}", text)

    def test_scan_template_text_odt(self):
        """scan_template_text 从 odt 提取文本"""
        tpl = self._path("t.odt")
        _create_minimal_odt(tpl, ["Hello {{name}}"])

        text = scan_template_text(tpl)
        self.assertIn("Hello {{name}}", text)

    def test_scan_template_text_unknown_raises(self):
        """scan_template_text 对未知格式抛 ValueError"""
        path = self._path("t.pdf")
        with open(path, "w") as f:
            f.write("not a doc")
        with self.assertRaises(ValueError):
            scan_template_text(path)


# =============================================================================
#  Renderer Registry Tests
# =============================================================================


class TestRendererRegistry(unittest.TestCase):
    """渲染器注册表测试"""

    def test_docx_registered_by_default(self):
        """DocxRenderer 默认注册"""
        cls = get_renderer("docx")
        self.assertIsNotNone(cls)
        self.assertEqual(cls, DocxRenderer)

    def test_odt_registered_by_default(self):
        """OdtRenderer 默认注册"""
        cls = get_renderer("odt")
        self.assertIsNotNone(cls)
        self.assertEqual(cls, OdtRenderer)

    def test_list_registered_formats(self):
        """list_registered_formats 返回所有已注册格式"""
        formats = list_registered_formats()
        self.assertIn("docx", formats)
        self.assertIn("odt", formats)

    def test_register_custom_renderer(self):
        """注册自定义渲染器"""
        class CustomRenderer(BaseRenderer):
            def load(self, path): pass
            def fill(self, data): pass
            def save(self, path): pass

        register_renderer("custom", CustomRenderer)
        self.assertEqual(get_renderer("custom"), CustomRenderer)
        self.assertIn("custom", list_registered_formats())


if __name__ == "__main__":
    unittest.main(verbosity=2)
