"""
DOCX 渲染器 — 基于 python-docx 的 Word 模板填充实现。

完全兼容现有 POCO 行为:
  - 正文段落 + 表格替换
  - 页眉 / 页脚替换
  - 跨 run 占位符安全处理
  - 残留占位符检测

依赖:
  - python-docx (>= 1.0.0)
"""

import os
import re
from typing import Dict, List

from docx import Document

from ..core.base_renderer import BaseRenderer

# 占位符正则：匹配 {{ ... }}
_PLACEHOLDER_RE = re.compile(r"\{\{(.*?)\}\}")


class DocxRenderer(BaseRenderer):
    """
    Microsoft Word (.docx) 模板渲染器。

    支持:
      - 正文段落 (paragraphs)
      - 表格 (tables) — 递归进入单元格
      - 页眉 (headers) — 所有节
      - 页脚 (footers) — 所有节
      - 跨 run 占位符安全替换（不删除图片/域代码/页码 run）

    Usage:
        renderer = DocxRenderer()
        renderer.load("template.docx")
        renderer.fill({"姓名": "PAN YU", "护照号": "EJ6376603"})
        renderer.save("output.docx")

        # 或一站式:
        renderer.process("template.docx", data, "output.docx")
    """

    ERROR_TAG = "DOCX_RENDER_ERROR"

    def __init__(self):
        self._doc: Document = None
        self._template_path: str = None

    # ---- BaseRenderer 接口实现 -----------------------------------------------

    def load(self, file_path: str) -> None:
        """
        加载 .docx 模板文件。

        Args:
            file_path: .docx 模板文件路径

        Raises:
            FileNotFoundError: 文件不存在
            ValueError:        文件不是有效的 .docx 格式
        """
        if not os.path.isfile(file_path):
            raise FileNotFoundError(
                self._error_tag(self.ERROR_TAG, f"模板文件不存在: {file_path}")
            )
        try:
            self._doc = Document(file_path)
            self._template_path = file_path
        except Exception as e:
            raise ValueError(
                self._error_tag(self.ERROR_TAG,
                                f"无法打开 .docx 文件: {file_path} — {e}")
            ) from e

    def fill(self, data: Dict[str, str]) -> None:
        """
        将数据填充到已加载的模板中。

        覆盖范围:
          - 正文段落 + 表格（递归）
          - 页眉 / 页脚（所有节）

        Args:
            data: {占位符名称: 替换值} 映射表
        """
        if self._doc is None:
            raise RuntimeError(
                self._error_tag(self.ERROR_TAG, "未加载模板，请先调用 load()")
            )

        # 1) 正文（段落 + 表格）
        self._replace_in_container(self._doc, data)

        # 2) 页眉 / 页脚
        for section in self._doc.sections:
            header = section.header
            if header:
                self._replace_in_container(header, data)

            footer = section.footer
            if footer:
                self._replace_in_container(footer, data)

    def save(self, output_path: str) -> None:
        """
        保存填充后的文档。

        Args:
            output_path: 输出 .docx 文件路径（目录不存在时自动创建）
        """
        if self._doc is None:
            raise RuntimeError(
                self._error_tag(self.ERROR_TAG, "未加载模板，请先调用 load()")
            )

        out_dir = os.path.dirname(output_path)
        if out_dir and not os.path.exists(out_dir):
            os.makedirs(out_dir, exist_ok=True)

        try:
            self._doc.save(output_path)
        except Exception as e:
            raise RuntimeError(
                self._error_tag(self.ERROR_TAG,
                                f"保存文档失败: {output_path} — {e}")
            ) from e

    # ---- 可选重写 ------------------------------------------------------------

    def extract_placeholders(self) -> List[str]:
        """
        从已加载的模板中提取所有占位符名称。

        扫描范围: 正文 + 表格 + 页眉 + 页脚

        Returns:
            占位符名称列表（去重，按首次出现顺序）
        """
        if self._doc is None:
            return []

        text = self._collect_all_text(self._doc)
        return self._parse_placeholders(text)

    def validate(self, data: Dict[str, str]) -> List[str]:
        """
        校验数据是否覆盖模板中的所有占位符。

        Returns:
            缺失的占位符名称列表
        """
        placeholders = self.extract_placeholders()
        missing = [ph for ph in placeholders if ph not in data]
        if missing:
            for ph in missing:
                print(
                    self._error_tag(self.ERROR_TAG,
                                    f"missing field: {ph}")
                )
        return missing

    # ---- 内部：段落替换逻辑 --------------------------------------------------

    def _replace_in_paragraph(self, paragraph, mapping: Dict[str, str]) -> None:
        """
        替换单个段落中所有 run 的占位符。

        两阶段:
          1) 简单替换 — 占位符完整包含在单个 run 内（覆盖 >95% 场景）
          2) 跨 run 替换 — 占位符跨多个 run 时的安全处理
        """
        if not paragraph.runs:
            return

        # ---- 阶段 1：逐 run 替换 ----
        for run in paragraph.runs:
            text = run.text
            if "{{" not in text:
                continue
            new_text = text
            for match in _PLACEHOLDER_RE.finditer(text):
                name = match.group(1).strip()
                if name in mapping:
                    new_text = new_text.replace(match.group(0), mapping[name])
            if new_text != text:
                run.text = new_text

        # ---- 阶段 2：跨 run 残留处理 ----
        remaining = paragraph.text
        if "{{" not in remaining and "}}" not in remaining:
            return

        self._replace_cross_run_placeholders(paragraph, mapping)

    def _replace_cross_run_placeholders(
        self, paragraph, mapping: Dict[str, str]
    ) -> None:
        """
        安全处理跨多个 run 的占位符。

        核心原则:
          - 只修改包含占位符文本的 run
          - 不删除任何 <w:r> 元素（图片、页码、域代码等不受影响）
          - 不在空文本 run（如图片 run）中添加 <w:t> 元素
          - 逐次处理右侧占位符，避免位置偏移
        """
        max_iterations = 50
        for _ in range(max_iterations):
            full_text = paragraph.text
            if "{{" not in full_text or "}}" not in full_text:
                return

            match = _PLACEHOLDER_RE.search(full_text)
            if not match:
                return

            name = match.group(1).strip()
            if name not in mapping:
                return

            match_start = match.start()
            match_end = match.end()
            replacement = mapping[name]

            runs = list(paragraph.runs)
            pos = 0
            involved = []

            for run in runs:
                run_start = pos
                run_end = pos + len(run.text)
                pos = run_end

                if run_end > match_start and run_start < match_end:
                    if len(run.text) > 0:
                        involved.append((run, run_start, run_end))

            if not involved:
                return

            if len(involved) == 1:
                run, rs, re_pos = involved[0]
                local_start = match_start - rs
                local_end = match_end - rs
                run.text = (
                    run.text[:local_start]
                    + replacement
                    + run.text[local_end:]
                )
            else:
                for i, (run, rs, re_pos) in enumerate(involved):
                    if i == 0:
                        prefix = run.text[: match_start - rs]
                        run.text = prefix + replacement
                    else:
                        local_end = match_end - rs
                        if local_end >= len(run.text):
                            run.text = ""
                        else:
                            run.text = run.text[local_end:]

    def _replace_in_container(self, container, mapping: Dict[str, str]) -> None:
        """
        在任意容器（document / header / footer / cell）中替换占位符。

        覆盖:
          - container.paragraphs
          - container.tables → 递归进入单元格
        """
        for paragraph in container.paragraphs:
            self._replace_in_paragraph(paragraph, mapping)

        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    self._replace_in_container(cell, mapping)

    # ---- 内部：文本收集（用于占位符提取）-------------------------------------

    def _collect_all_text(self, container) -> str:
        """递归收集容器中所有文本。"""
        parts = []

        for paragraph in container.paragraphs:
            parts.append(paragraph.text)

        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(self._collect_all_text(cell))

        # 页眉 / 页脚
        if hasattr(container, "sections"):
            for section in container.sections:
                header = section.header
                if header:
                    parts.append(self._collect_all_text(header))
                footer = section.footer
                if footer:
                    parts.append(self._collect_all_text(footer))

        return "".join(parts)

    @staticmethod
    def _parse_placeholders(text: str) -> List[str]:
        """从文本中提取占位符名称（去重，保留顺序）。"""
        seen = set()
        result: List[str] = []
        for match in _PLACEHOLDER_RE.finditer(text):
            name = match.group(1).strip()
            if name and name not in seen:
                seen.add(name)
                result.append(name)
        return result
