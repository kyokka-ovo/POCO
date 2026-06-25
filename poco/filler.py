"""
Word 模板填充引擎 — 将 {{占位符}} 替换为实际值并生成新文档。

替换范围：
  - 正文段落 (paragraphs)
  - 表格 (tables)
  - 页眉 (headers)
  - 页脚 (footers)

校验规则：
  - mapping 缺失字段 → 抛出 ValueError（rule_engine 必须覆盖全部字段）
  - 护照号/电话号等 → 格式校验 + warning
  - 输出文档残留 {{xxx}} → 抛出 ValueError
"""

import os
import re
from typing import Dict

from docx import Document

from .validator import validate_mapping, check_residual

# 占位符正则：匹配 {{ ... }}
_PLACEHOLDER_RE = re.compile(r"\{\{(.*?)\}\}")


# ---- 内部辅助函数 ---------------------------------------------------------


def _replace_in_paragraph(paragraph, mapping):
    """
    替换单个段落中所有 run 的占位符。

    分两阶段：
      1) 简单替换 —— 占位符完整包含在单个 run 内（覆盖 >95% 场景）
      2) 跨 run 替换 —— 占位符跨多个 run 时，只修改涉及的 run，
         绝不重建段落、绝不删除其他 run、绝不影响图片/页码/域代码
    """
    if not paragraph.runs:
        return

    # ---- 阶段 1：逐 run 替换 -----------------------------------------------
    for run in paragraph.runs:
        text = run.text
        if "{{" not in text:
            continue
        new_text = text
        for match in _PLACEHOLDER_RE.finditer(text):
            name = match.group(1).strip()
            if name in mapping:
                new_text = new_text.replace(match.group(0), mapping[name])
        if new_text != text:
            run.text = new_text

    # ---- 阶段 2：跨 run 残留处理 -------------------------------------------
    remaining = paragraph.text
    if "{{" not in remaining and "}}" not in remaining:
        return

    # 仍存在占位符 → 跨 run 安全替换（只改涉及的 run，不重建段落）
    _replace_cross_run_placeholders(paragraph, mapping)


def _replace_cross_run_placeholders(paragraph, mapping):
    """
    安全处理跨多个 run 的占位符。

    核心原则：
      - 只修改包含占位符文本的 run
      - 不删除任何 <w:r> 元素（图片、页码、域代码等不受影响）
      - 不在空文本 run（如图片 run）中添加 <w:t> 元素
      - 逐次处理右侧占位符，避免位置偏移

    策略：
      将替换值完整写入第一个涉及的 run，
      其余涉及的 run 仅清除其包含的占位符部分文本。
    """
    max_iterations = 50  # 安全上限，防止意外死循环
    for _ in range(max_iterations):
        full_text = paragraph.text
        if "{{" not in full_text or "}}" not in full_text:
            return

        # 查找第一个完整的 {{...}} 占位符
        match = _PLACEHOLDER_RE.search(full_text)
        if not match:
            return  # 无完整占位符（如仅有孤立的 {{ 或 }}）→ 安全退出

        name = match.group(1).strip()
        if name not in mapping:
            return  # 占位符不在 mapping 中 → 保留原样，停止处理

        match_start = match.start()
        match_end = match.end()
        replacement = mapping[name]

        # 定位涉及此占位符的所有 run
        runs = list(paragraph.runs)
        pos = 0
        involved = []  # [(run, run_start_char, run_end_char), ...]

        for run in runs:
            run_start = pos
            run_end = pos + len(run.text)
            pos = run_end

            if run_end > match_start and run_start < match_end:
                # 此 run 与占位符有交集
                if len(run.text) > 0:
                    # 仅包含有文本的 run（空文本 run 可能是图片、域代码等）
                    involved.append((run, run_start, run_end))

        if not involved:
            return

        if len(involved) == 1:
            # 单 run → 简单替换（阶段 1 应已处理；此处为防御）
            run, rs, re = involved[0]
            local_start = match_start - rs
            local_end = match_end - rs
            run.text = run.text[:local_start] + replacement + run.text[local_end:]
        else:
            # 多 run → 替换值完整写入第一个 run，其余 run 清除占位符部分
            for i, (run, rs, re) in enumerate(involved):
                if i == 0:
                    # 第一个涉及的 run：保留前缀 + 完整替换值
                    prefix = run.text[: match_start - rs]
                    run.text = prefix + replacement
                else:
                    # 其余涉及的 run：仅移除占位符覆盖的文本
                    local_end = match_end - rs
                    if local_end >= len(run.text):
                        run.text = ""
                    else:
                        run.text = run.text[local_end:]


def _replace_in_container(container, mapping):
    """
    在任意容器（document / header / footer / cell）中替换占位符。

    覆盖：
      - container.paragraphs
      - container.tables → 递归进入单元格
    """
    for paragraph in container.paragraphs:
        _replace_in_paragraph(paragraph, mapping)

    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_container(cell, mapping)


# ---- 公共 API -------------------------------------------------------------


def fill_template(
    template_path: str,
    output_path: str,
    mapping: Dict[str, str],
) -> None:
    """
    读取 Word 模板，校验 + 替换所有 {{占位符}}，保存到 output_path。

    执行流程：
      1) validate_mapping()   → 检测缺失字段 / 格式错误
      2) 缺失字段 → 抛出 ValueError（rule_engine 必须覆盖全部字段）
      3) 文档填充              → 正文 / 表格 / 页眉 / 页脚
      4) check_residual()     → 残留占位符检测（有则抛 ValueError）

    Args:
        template_path: 输入 .docx 模板文件路径
        output_path:  输出 .docx 文件路径（目录不存在时自动创建）
        mapping:       {占位符名称: 替换值} 映射表

    Replacement scope:
        - 正文段落 (paragraphs)
        - 表格 (tables)
        - 页眉 (headers)
        - 页脚 (footers)

    Guarantees:
        - 输出 100% 无 {{xxx}} 残留
        - 所有字段均由 rule_engine 生成，不存在 [未填写:xxx] 兜底

    Example:
        >>> fill_template(
        ...     "template.docx",
        ...     "output/generated.docx",
        ...     {"姓名": "PAN YU", "护照号": "EJ6376603"},
        ... )
    """
    # ---- 阶段 1：校验 ----------------------------------------------------
    result = validate_mapping(template_path, mapping)

    # 输出校验报告
    report = result.report()
    if result.has_warnings:
        print(report, flush=True)

    # 缺失字段视为硬错误：rule_engine 必须覆盖全部字段
    if result.missing_fields:
        raise ValueError(
            f"Mapping 缺失字段（{len(result.missing_fields)} 个）："
            + ", ".join(result.missing_fields)
        )

    # ---- 阶段 2：填充 ----------------------------------------------------
    _fill_document(template_path, output_path, mapping)

    # ---- 阶段 3：残留检测 ------------------------------------------------
    residual = check_residual(output_path)
    if residual:
        raise ValueError(
            f"未完全填充占位符：{', '.join(residual)}"
        )


def _fill_document(
    template_path: str,
    output_path: str,
    mapping: Dict[str, str],
) -> None:
    """
    内部：纯填充逻辑（不做校验）。

    直接读取模板、替换占位符、写入输出。
    """
    # 自动创建输出目录
    out_dir = os.path.dirname(output_path)
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir, exist_ok=True)

    doc = Document(template_path)

    # 1) 正文（段落 + 表格）
    _replace_in_container(doc, mapping)

    # 2) 页眉 / 页脚
    for section in doc.sections:
        header = section.header
        if header:
            _replace_in_container(header, mapping)

        footer = section.footer
        if footer:
            _replace_in_container(footer, mapping)

    doc.save(output_path)
